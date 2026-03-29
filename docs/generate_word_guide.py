"""
Pulse User's Guide — Word Document Generator
Generates a print-ready .docx with cover sheet (Pulse logo) + full guide content.
"""

import re
import io
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Paths ────────────────────────────────────────────────────────────────────
SCRIPT_DIR = Path(__file__).parent
LOGO_PATH  = Path(r"C:\Users\Aegis{FM}\Downloads\pulse-logo-helix.png")
GUIDE_MD   = SCRIPT_DIR / "USER_GUIDE.md"
OUTPUT     = SCRIPT_DIR / "Pulse_Users_Guide_v25.1.3.docx"

# ── Brand colours ─────────────────────────────────────────────────────────────
ROSE      = RGBColor(0xF4, 0x3F, 0x5E)   # #f43f5e
DARK_NAVY = RGBColor(0x0F, 0x17, 0x2A)   # #0f172a
DARK_BG   = RGBColor(0x18, 0x1F, 0x36)   # #181f36
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
GRAY_700  = RGBColor(0x37, 0x41, 0x51)
GRAY_500  = RGBColor(0x6B, 0x72, 0x80)
GRAY_200  = RGBColor(0xE5, 0xE7, 0xEB)
ROSE_LIGHT= RGBColor(0xFF, 0xE4, 0xE6)   # #ffe4e6
ROSE_DARK = RGBColor(0x9F, 0x12, 0x39)   # #9f1239

# ── Helper: set paragraph background shading ─────────────────────────────────
def shade_paragraph(paragraph, hex_fill: str):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_fill)
    pPr.append(shd)

def shade_cell(cell, hex_fill: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_fill)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    """Set borders on a table cell. kwargs: top, bottom, left, right — each a dict."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        tag = 'w:' + edge
        setting = kwargs.get(edge, {})
        if setting:
            el = OxmlElement(tag)
            el.set(qn('w:val'),   setting.get('val',   'single'))
            el.set(qn('w:sz'),    setting.get('sz',    '4'))
            el.set(qn('w:color'), setting.get('color', 'auto'))
            tcBorders.append(el)
    tcPr.append(tcBorders)

def page_break(doc):
    doc.add_page_break()

def add_run_with_style(para, text, bold=False, italic=False, color=None, size=None, font_name=None):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    if font_name:
        run.font.name = font_name
    return run

# ── Build document styles ─────────────────────────────────────────────────────
def setup_styles(doc):
    styles = doc.styles

    # Normal (body)
    normal = styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = GRAY_700
    normal.paragraph_format.space_after  = Pt(6)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.line_spacing = Pt(15)

    # Heading 1
    h1 = styles['Heading 1']
    h1.font.name  = 'Calibri'
    h1.font.size  = Pt(20)
    h1.font.bold  = True
    h1.font.color.rgb = ROSE
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after  = Pt(8)
    h1.paragraph_format.keep_with_next = True

    # Heading 2
    h2 = styles['Heading 2']
    h2.font.name  = 'Calibri'
    h2.font.size  = Pt(14)
    h2.font.bold  = True
    h2.font.color.rgb = DARK_NAVY
    h2.paragraph_format.space_before = Pt(16)
    h2.paragraph_format.space_after  = Pt(4)
    h2.paragraph_format.keep_with_next = True

    # Heading 3
    h3 = styles['Heading 3']
    h3.font.name  = 'Calibri'
    h3.font.size  = Pt(11.5)
    h3.font.bold  = True
    h3.font.color.rgb = ROSE
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after  = Pt(2)
    h3.paragraph_format.keep_with_next = True

    # List Bullet
    try:
        lb = styles['List Bullet']
    except KeyError:
        lb = styles.add_style('List Bullet', WD_STYLE_TYPE.PARAGRAPH)
    lb.font.name = 'Calibri'
    lb.font.size = Pt(10.5)
    lb.paragraph_format.left_indent   = Inches(0.3)
    lb.paragraph_format.space_after   = Pt(3)

    # List Number
    try:
        ln = styles['List Number']
    except KeyError:
        ln = styles.add_style('List Number', WD_STYLE_TYPE.PARAGRAPH)
    ln.font.name = 'Calibri'
    ln.font.size = Pt(10.5)
    ln.paragraph_format.left_indent   = Inches(0.3)
    ln.paragraph_format.space_after   = Pt(3)

    # Tip box style (we'll fake it with shading)
    try:
        tip_style = styles.add_style('Tip Box', WD_STYLE_TYPE.PARAGRAPH)
    except:
        tip_style = styles['Normal']
    tip_style.font.name = 'Calibri'
    tip_style.font.size = Pt(10)
    tip_style.font.italic = True
    tip_style.paragraph_format.left_indent  = Inches(0.3)
    tip_style.paragraph_format.space_after  = Pt(6)
    tip_style.paragraph_format.space_before = Pt(6)

    return doc

# ── Cover Page ────────────────────────────────────────────────────────────────
def build_cover(doc, logo_path: Path):
    """Full-page black cover — logo centered, nothing else."""
    section = doc.sections[0]
    section.page_width    = Inches(8.5)
    section.page_height   = Inches(11)
    section.top_margin    = Inches(0)
    section.bottom_margin = Inches(0)
    section.left_margin   = Inches(0)
    section.right_margin  = Inches(0)

    # Single full-page black cell — logo centered inside
    cover_table = doc.add_table(rows=1, cols=1)
    cover_table.style = 'Table Grid'
    cell = cover_table.cell(0, 0)
    shade_cell(cell, '000000')
    for edge in ('top', 'left', 'bottom', 'right'):
        set_cell_border(cell, **{edge: {'val': 'none', 'sz': '0', 'color': '000000'}})
    cover_table.rows[0].height = Inches(11)

    # Vertical centering via top spacer paragraph
    # Page is 11" — logo at ~3.8" wide, position vertically centred ≈ (11 - logo_h) / 2
    # We achieve rough vertical centre with space_before on the image paragraph.
    logo_para = cell.paragraphs[0]
    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_para.paragraph_format.space_before = Pt(180)   # ~2.5" push-down
    logo_para.paragraph_format.space_after  = Pt(0)

    run = logo_para.add_run()
    if logo_path.exists():
        run.add_picture(str(logo_path), width=Inches(4.2))
    else:
        run.text = 'PULSE'
        run.bold = True
        run.font.size = Pt(72)
        run.font.color.rgb = ROSE

    page_break(doc)

    # Reset margins for body pages
    body_section = doc.add_section()
    body_section.page_width    = Inches(8.5)
    body_section.page_height   = Inches(11)
    body_section.top_margin    = Inches(1.0)
    body_section.bottom_margin = Inches(1.0)
    body_section.left_margin   = Inches(1.2)
    body_section.right_margin  = Inches(1.2)

# ── Markdown parser → Word ────────────────────────────────────────────────────

def strip_md_inline(text: str) -> str:
    """Remove inline markdown: **bold**, *italic*, `code`, [text](url)."""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*',     r'\1', text)
    text = re.sub(r'`(.+?)`',       r'\1', text)
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    return text

def add_inline_md(para, text: str, base_color=None):
    """Add text to paragraph preserving **bold** and *italic* inline."""
    base_color = base_color or GRAY_700
    # tokenise: **bold**, *italic*, `code`, plain
    pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)'
    parts = re.split(pattern, text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            r = para.add_run(part[2:-2])
            r.bold = True
            r.font.color.rgb = base_color
            r.font.name = 'Calibri'
        elif part.startswith('*') and part.endswith('*'):
            r = para.add_run(part[1:-1])
            r.italic = True
            r.font.color.rgb = base_color
            r.font.name = 'Calibri'
        elif part.startswith('`') and part.endswith('`'):
            r = para.add_run(part[1:-1])
            r.font.name = 'Courier New'
            r.font.size = Pt(9.5)
            r.font.color.rgb = ROSE
        else:
            if part:
                r = para.add_run(part)
                r.font.color.rgb = base_color
                r.font.name = 'Calibri'

def parse_md_to_doc(doc, md_text: str):
    lines = md_text.split('\n')
    i = 0
    in_table = False
    table_rows = []
    in_details = False  # skip <details> blocks body but show summary as tip

    while i < len(lines):
        line = lines[i]
        raw = line

        # ── HTML: skip <details> open/close, convert <summary> → tip-style box
        if re.match(r'\s*<details', line, re.I):
            in_details = True
            i += 1
            continue
        if re.match(r'\s*</details>', line, re.I):
            in_details = False
            i += 1
            continue
        m_summary = re.match(r'\s*<summary[^>]*>(.*?)</summary>', line, re.I)
        if m_summary:
            # Extract text from summary (strip inner HTML)
            text = re.sub(r'<[^>]+>', '', m_summary.group(1)).strip()
            p = doc.add_paragraph(style='Normal')
            shade_paragraph(p, 'FFF1F2')  # rose-50
            p.paragraph_format.left_indent  = Inches(0.2)
            p.paragraph_format.right_indent = Inches(0.2)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after  = Pt(4)
            r = p.add_run('▶  ' + text)
            r.bold = True
            r.font.color.rgb = ROSE_DARK
            r.font.name = 'Calibri'
            r.font.size = Pt(10)
            i += 1
            continue

        # ── Table ──────────────────────────────────────────────────────────────
        if '|' in line and line.strip().startswith('|'):
            cells = [c.strip() for c in line.strip().strip('|').split('|')]
            # separator row?
            if all(re.match(r'^[-: ]+$', c) for c in cells if c):
                i += 1
                continue
            table_rows.append(cells)
            i += 1
            # peek ahead — if next line is not a table row, flush
            if i >= len(lines) or not (lines[i].strip().startswith('|') and '|' in lines[i]):
                if table_rows:
                    col_count = max(len(r) for r in table_rows)
                    # ensure all rows have same col count
                    table_rows = [r + [''] * (col_count - len(r)) for r in table_rows]
                    tbl = doc.add_table(rows=len(table_rows), cols=col_count)
                    tbl.style = 'Table Grid'
                    for ri, row_data in enumerate(table_rows):
                        for ci, cell_text in enumerate(row_data):
                            cell = tbl.cell(ri, ci)
                            cell_text_clean = strip_md_inline(cell_text)
                            p = cell.paragraphs[0]
                            p.paragraph_format.space_before = Pt(3)
                            p.paragraph_format.space_after  = Pt(3)
                            add_inline_md(p, cell_text, DARK_NAVY if ri == 0 else GRAY_700)
                            if ri == 0:
                                shade_cell(cell, 'F43F5E')
                                for r in p.runs:
                                    r.font.color.rgb = WHITE
                                    r.bold = True
                            elif ri % 2 == 0:
                                shade_cell(cell, 'FFF1F2')
                            else:
                                shade_cell(cell, 'FFFFFF')
                    doc.add_paragraph('')
                    table_rows = []
            continue

        # ── Headings ──────────────────────────────────────────────────────────
        if line.startswith('#### '):
            text = line[5:].strip()
            p = doc.add_paragraph(style='Heading 3')
            add_inline_md(p, text, DARK_NAVY)
            i += 1
            continue
        if line.startswith('### '):
            text = line[4:].strip()
            p = doc.add_paragraph(style='Heading 3')
            add_inline_md(p, text, ROSE)
            i += 1
            continue
        if line.startswith('## '):
            text = line[3:].strip()
            p = doc.add_paragraph(style='Heading 2')
            add_inline_md(p, text, DARK_NAVY)
            i += 1
            continue
        if line.startswith('# '):
            text = line[2:].strip()
            p = doc.add_paragraph(style='Heading 1')
            add_inline_md(p, text, ROSE)
            i += 1
            continue

        # ── Horizontal rule ───────────────────────────────────────────────────
        if re.match(r'^---+\s*$', line):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'),   'single')
            bottom.set(qn('w:sz'),    '6')
            bottom.set(qn('w:color'), 'F43F5E')
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # ── Blockquote / warning ──────────────────────────────────────────────
        if line.startswith('> '):
            text = line[2:].strip()
            p = doc.add_paragraph(style='Normal')
            shade_paragraph(p, 'FFF7ED')  # orange-50
            p.paragraph_format.left_indent  = Inches(0.3)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            add_inline_md(p, '⚠  ' + text, RGBColor(0xC2, 0x41, 0x0C))
            i += 1
            continue

        # ── Tip: line ─────────────────────────────────────────────────────────
        if line.strip().startswith('**Tip:**') or line.strip().startswith('**Tip: **'):
            p = doc.add_paragraph(style='Normal')
            shade_paragraph(p, 'EFF6FF')  # blue-50
            p.paragraph_format.left_indent  = Inches(0.25)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(6)
            add_inline_md(p, '💡  ' + line.strip(), RGBColor(0x1D, 0x4E, 0xD8))
            i += 1
            continue

        # ── Tip: note: ────────────────────────────────────────────────────────
        if line.strip().startswith('**Note:**'):
            p = doc.add_paragraph(style='Normal')
            shade_paragraph(p, 'F0FDF4')
            p.paragraph_format.left_indent  = Inches(0.25)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(6)
            add_inline_md(p, '📝  ' + line.strip(), RGBColor(0x16, 0x65, 0x34))
            i += 1
            continue

        # ── Code block ───────────────────────────────────────────────────────
        if line.startswith('```'):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            # Skip code blocks (not user-guide content)
            continue

        # ── Ordered list ──────────────────────────────────────────────────────
        m_ol = re.match(r'^(\s*)(\d+)\.\s+(.*)', line)
        if m_ol:
            indent = len(m_ol.group(1)) // 3
            text = m_ol.group(3).strip()
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.left_indent   = Inches(0.3 + indent * 0.25)
            p.paragraph_format.first_line_indent = Inches(-0.2)
            p.paragraph_format.space_after = Pt(2)
            add_inline_md(p, text, GRAY_700)
            i += 1
            continue

        # ── Unordered list ────────────────────────────────────────────────────
        m_ul = re.match(r'^(\s*)[-*]\s+(.*)', line)
        if m_ul:
            indent = len(m_ul.group(1)) // 2
            text = m_ul.group(2).strip()
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent   = Inches(0.3 + indent * 0.2)
            p.paragraph_format.first_line_indent = Inches(-0.15)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run('• ')
            r.font.color.rgb = ROSE
            r.font.bold = True
            add_inline_md(p, text, GRAY_700)
            i += 1
            continue

        # ── Blank line ────────────────────────────────────────────────────────
        if line.strip() == '':
            i += 1
            continue

        # ── Normal paragraph ──────────────────────────────────────────────────
        p = doc.add_paragraph(style='Normal')
        add_inline_md(p, line.strip(), GRAY_700)
        i += 1

# ── Main ─────────────────────────────────────────────────────────────────────

def main():
    print("Reading USER_GUIDE.md …")
    md_text = GUIDE_MD.read_text(encoding='utf-8')

    # Strip the metadata header (lines 1–6) since we put it on the cover
    lines = md_text.split('\n')
    # Find first `---` separator and skip through it
    body_start = 0
    dashes_seen = 0
    for idx, l in enumerate(lines):
        if re.match(r'^---+\s*$', l):
            dashes_seen += 1
            if dashes_seen == 1:
                body_start = idx + 1
                break
    md_body = '\n'.join(lines[body_start:])

    print("Building Word document …")
    doc = Document()
    setup_styles(doc)

    # ── Cover ─────────────────────────────────────────────────────────────────
    build_cover(doc, LOGO_PATH)

    # ── Body ──────────────────────────────────────────────────────────────────
    parse_md_to_doc(doc, md_body)

    # ── Footer on body pages ─────────────────────────────────────────────────
    for section in doc.sections[1:]:
        footer = section.footer
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = fp.add_run('Pulse User\'s Guide  ·  Version 25.1.3  ·  March 17, 2026  ·  pulse.app')
        r1.font.size = Pt(8)
        r1.font.color.rgb = GRAY_500
        r1.font.name = 'Calibri'

    print(f"Saving to {OUTPUT} …")
    doc.save(str(OUTPUT))
    print(f"Done! Output:  {OUTPUT}")

if __name__ == '__main__':
    main()
